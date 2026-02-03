"""文档解析器

支持多格式文档解析：txt, md, pdf, docx
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from app.core.logging import get_logger

logger = get_logger("document_parser")


@dataclass
class Section:
    """文档章节"""
    level: int              # 标题级别 (1-6)
    title: str              # 标题文本
    content: str            # 章节内容
    start_pos: int = 0      # 开始位置
    end_pos: int = 0        # 结束位置


@dataclass
class ParsedDocument:
    """解析后的文档"""
    content: str                          # 纯文本内容
    title: str | None = None              # 提取的标题
    sections: list[Section] = field(default_factory=list)  # 章节结构
    metadata: dict[str, Any] = field(default_factory=dict) # 元数据
    file_type: str = ""                   # 文件类型
    word_count: int = 0                   # 字数
    char_count: int = 0                   # 字符数


class UnsupportedFormatError(Exception):
    """不支持的文件格式"""
    pass


class DocumentParser:
    """文档解析器"""

    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
    }

    async def parse(self, file_path: str) -> ParsedDocument:
        """解析文档

        Args:
            file_path: 文件路径

        Returns:
            ParsedDocument

        Raises:
            UnsupportedFormatError: 不支持的文件格式
            FileNotFoundError: 文件不存在
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(f"不支持的格式: {ext}")

        file_type = self.SUPPORTED_EXTENSIONS[ext]

        logger.info(f"解析文档: {file_path}, 类型: {file_type}")

        if file_type == 'text':
            return await self._parse_text(path)
        elif file_type == 'markdown':
            return await self._parse_markdown(path)
        elif file_type == 'pdf':
            return await self._parse_pdf(path)
        elif file_type in ['docx', 'doc']:
            return await self._parse_docx(path)
        else:
            raise UnsupportedFormatError(f"解析器未实现: {file_type}")

    async def _parse_text(self, path: Path) -> ParsedDocument:
        """解析纯文本文件"""
        content = path.read_text(encoding='utf-8')
        
        # 尝试从第一行提取标题
        lines = content.split('\n')
        title = lines[0].strip() if lines else None
        
        return ParsedDocument(
            content=content,
            title=title,
            file_type='text',
            word_count=len(content.split()),
            char_count=len(content),
            metadata={
                'file_size': path.stat().st_size,
                'line_count': len(lines),
            },
        )

    async def _parse_markdown(self, path: Path) -> ParsedDocument:
        """解析 Markdown 文件"""
        content = path.read_text(encoding='utf-8')
        
        # 提取标题（第一个 # 标题）
        title = None
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if title_match:
            title = title_match.group(1).strip()
        
        # 提取章节结构
        sections = self._extract_markdown_sections(content)
        
        # 提取代码块
        code_blocks = re.findall(r'```[\w]*\n(.*?)```', content, re.DOTALL)
        
        return ParsedDocument(
            content=content,
            title=title,
            sections=sections,
            file_type='markdown',
            word_count=len(content.split()),
            char_count=len(content),
            metadata={
                'file_size': path.stat().st_size,
                'code_block_count': len(code_blocks),
                'section_count': len(sections),
            },
        )

    def _extract_markdown_sections(self, content: str) -> list[Section]:
        """提取 Markdown 章节结构"""
        sections = []
        pattern = r'^(#{1,6})\s+(.+)$'
        
        matches = list(re.finditer(pattern, content, re.MULTILINE))
        
        for i, match in enumerate(matches):
            level = len(match.group(1))
            title = match.group(2).strip()
            start_pos = match.end()
            
            # 找到下一个同级或更高级标题
            end_pos = len(content)
            for next_match in matches[i + 1:]:
                next_level = len(next_match.group(1))
                if next_level <= level:
                    end_pos = next_match.start()
                    break
            
            section_content = content[start_pos:end_pos].strip()
            
            sections.append(Section(
                level=level,
                title=title,
                content=section_content,
                start_pos=match.start(),
                end_pos=end_pos,
            ))
        
        return sections

    async def _parse_pdf(self, path: Path) -> ParsedDocument:
        """解析 PDF 文件"""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber 未安装，尝试使用 PyPDF2")
            try:
                import pypdf
                return await self._parse_pdf_with_pypdf(path)
            except ImportError:
                raise UnsupportedFormatError(
                    "PDF 解析需要安装 pdfplumber 或 pypdf: pip install pdfplumber"
                )

        content_parts = []
        page_count = 0
        
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content_parts.append(text)
        
        content = '\n\n'.join(content_parts)
        
        # 尝试从第一页提取标题
        title = None
        if content_parts:
            first_lines = content_parts[0].split('\n')[:3]
            for line in first_lines:
                if len(line.strip()) > 5:
                    title = line.strip()
                    break
        
        return ParsedDocument(
            content=content,
            title=title,
            file_type='pdf',
            word_count=len(content.split()),
            char_count=len(content),
            metadata={
                'file_size': path.stat().st_size,
                'page_count': page_count,
            },
        )

    async def _parse_pdf_with_pypdf(self, path: Path) -> ParsedDocument:
        """使用 pypdf 解析 PDF"""
        import pypdf
        
        content_parts = []
        
        with open(path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            page_count = len(reader.pages)
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content_parts.append(text)
        
        content = '\n\n'.join(content_parts)
        
        title = None
        if content_parts:
            first_lines = content_parts[0].split('\n')[:3]
            for line in first_lines:
                if len(line.strip()) > 5:
                    title = line.strip()
                    break
        
        return ParsedDocument(
            content=content,
            title=title,
            file_type='pdf',
            word_count=len(content.split()),
            char_count=len(content),
            metadata={
                'file_size': path.stat().st_size,
                'page_count': page_count,
            },
        )

    async def _parse_docx(self, path: Path) -> ParsedDocument:
        """解析 Word 文档"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise UnsupportedFormatError(
                "Word 解析需要安装 python-docx: pip install python-docx"
            )

        doc = DocxDocument(path)
        
        paragraphs = []
        title = None
        sections = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            paragraphs.append(text)
            
            # 检测标题样式
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if 'heading' in style_name or 'title' in style_name:
                    # 提取标题级别
                    level = 1
                    level_match = re.search(r'(\d)', style_name)
                    if level_match:
                        level = int(level_match.group(1))
                    
                    if title is None and level == 1:
                        title = text
                    
                    sections.append(Section(
                        level=level,
                        title=text,
                        content="",  # 将在后续填充
                    ))
        
        content = '\n\n'.join(paragraphs)
        
        # 如果没有从样式检测到标题，使用第一段
        if title is None and paragraphs:
            title = paragraphs[0]
        
        return ParsedDocument(
            content=content,
            title=title,
            sections=sections,
            file_type='docx',
            word_count=len(content.split()),
            char_count=len(content),
            metadata={
                'file_size': path.stat().st_size,
                'paragraph_count': len(paragraphs),
            },
        )

    def get_supported_extensions(self) -> list[str]:
        """获取支持的文件扩展名"""
        return list(self.SUPPORTED_EXTENSIONS.keys())
